from os.path import exists

import numpy as np
import pandas as pd
from tqdm import tqdm
from termcolor import colored
import random
import copy
import src.tasks.visual_results
from docx import Document

from src.models.nbhd import NBHD

import sys


class Task:

    def __init__(self, task_name, env, bb_model, dataset, method, method_name, search_objs, eval_objs, eval_path, visual_path, nbhd):
        self.task_name = task_name
        self.env = env
        self.bb_model = bb_model
        self.dataset = dataset
        self.method = method
        self.method_name = method_name
        self.eval_path = eval_path
        self.visual_path = visual_path
        self.search_objs = search_objs
        self.eval_objs = eval_objs
        self.nbhd = nbhd
        self.ACTIONS = {0: 'SOUTH', 1: 'NORTH', 2: 'EAST', 3: 'WEST', 4: 'PICKUP',5: 'DROPOFF'}
        self.LOCS = {0: 'Red', 1: 'Green', 2: 'Yellow', 3: 'Blue', 4: 'In Taxi'}

    def run_experiment(self, cf_typename, facts, targets=None):
        print('Running experiment for {} task with {}'.format(self.task_name, self.method_name))
        print('Finding counterfactuals for {} facts'.format(len(facts)))
        doc = Document()
        
        # get cfs for facts
        eval_dict = {}
        cnt = 0
        nonecount = 0
        foundcount = 0
        wrongcount = []
        evalavgs = {'cost': 0.0, 'reachability': 0.0, 'stochasticity': 0.0, 'validity': 0.0, 'realistic': 0.0, 'actionable': 0.0, 'proximity': 0.0, 'sparsity': 0.0, 'dmc': 0.0}
        for i in tqdm(range(len(facts))):
            f = facts[i]

            if isinstance(f, dict):
                f = self.env.generate_state_from_json(f)

            if self.nbhd:
                nbhd = NBHD(self.env, f, max_level=0)
            else:
                nbhd = None

            if targets is None:
                ts = self.get_targets(f, self.env, self.bb_model)
            else:
                ts = [targets[i]]
            illgealAll = 0
            for t in ts:
                print('FACT: Target = {}'.format(self.ACTIONS[t]), ': State = ' , self.bb_model.env.encode(f[0], f[1], f[2], f[3]), f)
                self.env.render_state(f)
                cf = self.method.generate_counterfactuals(f, t, nbhd)
                statement = ""
                # wrongcount += wc

                if cf is None:
                    nonecount += 1
                    found = False
                    self.evaluate_cf(f, t, cf, found)
                    #if(not(self.env.check_done(f)) ):
                        #and not(list(f) == list([0,2,2,0]))
                        #sys.exit()
                    doc = src.tasks.visual_results.save_visual_results(self.env, doc, f, cf, t, statement, cf_typename)
                    doc.save(self.visual_path)
                    wrongcount = []
                    # for i in range(self.env.max_row+1):
                    #     for j in range(self.env.max_col+1):
                    #         check = [i,j,f[2],f[3]]
                    #         if(self.bb_model.predict(check) == t):
                    #             wrongcount += ["Fact: " + str(f) + " found target: " + str(t) + " at " + str(check)]
   

                else:
                    foundcount += 1
                    found = True
                    predicted_action = self.ACTIONS[self.bb_model.predict(f)]
                    encoded_fact = self.bb_model.env.encode(f[0], f[1], f[2], f[3])

                    
                    rews = self.evaluate_cf(f, t, cf, found)
                    for key, value in rews.items():
                        if key in evalavgs:
                            evalavgs[key] += value[0]
                        else:
                            evalavgs[key] = value

                    print("State Path: ", cf.path)
                    #print("Loss function value:", cf.value)
                    print("Action path: ", cf.action_path)
                    if(cf_typename == 'ACTION' or cf_typename == 'ACTION_CHANGELOC'):
                        if(f[2] != cf.cf_state[2]):
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) \
                            + ' where the passanger location is now ' + self.LOCS[cf.cf_state[2]]+ ':'
                            print(statement)
                        else:
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) + ':'
                            print(statement)
                        
                        self.env.render_state(cf.cf_state)
                    elif(cf_typename == 'ACTION_CHANGEDST'):
                        if(f[3] != cf.cf_state[3]):
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) \
                            + ' where the destination is now ' + self.LOCS[cf.cf_state[3]]+ ':'
                            print(statement)
                        else:
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) + ':'
                            print(statement)
                        
                        self.env.render_state(cf.cf_state)
                    if(cf_typename == 'ACTION_CHANGEBOTH'):
                        print("here", f, cf.cf_state)
                        if(f[2] != cf.cf_state[2]):
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) \
                            + ' where the passanger location is now ' + self.LOCS[cf.cf_state[2]]+ ':'
                            print(statement)
                        elif(f[3] != cf.cf_state[3]):
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) \
                            + ' where the destination is now ' + self.LOCS[cf.cf_state[3]]+ ':'
                            print(statement)
                        else:
                            statement = 'Statement: Given that in state ' + str(encoded_fact) \
                            + ' the agent chose ' + predicted_action \
                            + ' in what state would it chose '+ str(self.ACTIONS[t]) \
                            + '? \nIn state ' + str(self.bb_model.env.encode(cf.cf_state[0], cf.cf_state[1], cf.cf_state[2], cf.cf_state[3])) + ':'
                            print(statement)
                        
                        self.env.render_state(cf.cf_state)
                    elif(cf_typename == 'NETACTION'):
                        print(cf.cf_state)
                        # news, r, d, _ = self.env.step(self.bb_model.predict(cf.cf_state))
                        # cf.cf_state = news
                        rowcol = ""
                        rowcolnum1 = -1
                        rowcolnum2 = -1
                        
                        # cancels = {'EAST': 'WEST', 'WEST': 'EAST', 'NORTH': 'SOUTH', 'SOUTH': 'NORTH'}
                        # canceled_path = []

                        # for direction in cf.action_path:
                        #     opposite = cancels.get(direction)
                        #     if canceled_path and canceled_path[-1] == opposite:
                        #         canceled_path.pop()
                        #     else:
                        #         canceled_path.append(direction)

                        if(predicted_action == "NORTH" or predicted_action == "SOUTH"):
                             rowcol = "column"
                             rowcolnum1 = f[1]
                             rowcolnum2 = cf.cf_state[1]
                        else:
                             rowcol = "row"
                             rowcolnum1 = f[0]
                             rowcolnum2 = cf.cf_state[0]
                        #     rowcolnum1 = f[0]
                        # else:
                        #     rowcol1 = "column"
                        #     rowcolnum1 = f[1]
                        # if(t == 0 or t == 1):
                        #     rowcol2 = "row"
                        #     numofdirection = canceled_path.count(self.ACTIONS[t])
                        #     if(self.ACTIONS[t] == "SOUTH"):
                        #         rowcolnum2 = rowcolnum1 + numofdirection
                        #     else:
                        #         rowcolnum2 = rowcolnum1 - numofdirection
                        

                        #     numofdirection = canceled_path.count(self.ACTIONS[t])
                        #     if(self.ACTIONS[t] == "EAST"):
                        #         rowcolnum2 = rowcolnum1 + numofdirection
                        #     else:
                        #         rowcolnum2 = rowcolnum1 - numofdirection

                        statement = 'Statement: Given that in state ' + str(encoded_fact) \
                        + ' the agent chose to travel ' + predicted_action \
                        + ' on ' + str(rowcol) + ' ' + str(rowcolnum1) \
                        + ' in what state would it have travelled ' + predicted_action \
                        + ' on ' + str(rowcol) + ' ' + str(rowcolnum2) + ' instead?' \
                        + '\n  In state ' + str(cf.cf_state)#str(self.env.encode(news[0], news[1], news[2], news[3])) +' :'
                        print(statement)
                        self.env.render_state(cf.cf_state)
                        #cf.cf_state = news
                
                doc = src.tasks.visual_results.save_visual_results(self.env, doc, f, cf, t, statement, cf_typename)
                doc.save(self.visual_path)
                cnt += 1
            print("\n\n")

        accuracy = foundcount/(nonecount+foundcount)
        print("CF Accuracy: ", accuracy)
        for key, value in evalavgs.items():
            evalavgs[key] = evalavgs[key]/foundcount
        print("Eval avgs: ", evalavgs)
        doc.add_paragraph().add_run("CF Accuracy: " + str(accuracy) + 'in mode ' + cf_typename)
        print("LBO = ", evalavgs['proximity'] + evalavgs['sparsity'] + evalavgs['dmc'])
        print("L = ", evalavgs['reachability'] + evalavgs['cost'] + evalavgs['stochasticity'])
        doc.save(self.visual_path)
        wrongcount = list(set(wrongcount))
        print('WRONG:')
        for i in wrongcount:
            print(i)

        

    def get_targets(self, f, env, bb_model):
        pred = bb_model.predict(f)
        available_actions = env.get_actions(f)
        targets = [a for a in available_actions if a != pred]
        print('Alternate Actions: ', targets, 'Predicted: ', pred)

        return targets

    def evaluate_cf(self, f, t, cf, found):
        if not found:
            eval_obj_names = []

            for obj in self.eval_objs:
                eval_obj_names += list(obj.lmbdas.keys())

            ind_rew = [0] * len(eval_obj_names)
            df = pd.DataFrame([ind_rew], columns=eval_obj_names)

            df['total_reward'] = 0
            df['cf'] = 0

        else:
            rews = {}
            for obj in self.eval_objs:
                ind_rew, total_rew = obj.get_ind_rews(f, cf.cf_state,  t, cf.actions, cf.cumulative_reward)
                ind_rew = {k: [v] for k, v in ind_rew.items()}
                rews.update(ind_rew)

            df = pd.DataFrame([rews])
            total_rew = self.search_objs.get_reward(f, cf.cf_state, t, cf.actions, cf.cumulative_reward)

            print(rews)
            ind_rew = copy.deepcopy(rews)
            df['total_reward'] = total_rew

            df['cf'] = self.env.cf_writable_state(cf.cf_state, cf.action_path)

        df['fact'] = list(np.tile(self.env.writable_state(f), (len(df), 1)))
        df['target'] = t
        df['found'] = found

        header = not exists(self.eval_path)
        df.to_csv(self.eval_path, mode='a', header=header)

        return ind_rew

    def true_path(self, env, model, start_state):
        env.reset()
        env.set_state(start_state)

        done = False
        true_path = []
        obs = start_state
        while not done:
            action = model.predict(obs)
            true_path.append(action)
            obs, rew, done, _ = env.step(action)

        return true_path

