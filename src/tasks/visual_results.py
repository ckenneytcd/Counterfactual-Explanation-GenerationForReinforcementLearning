from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK
from colorama import init, Fore, Back, Style
from docx.oxml.ns import qn

ACTIONS = {0: 'SOUTH', 1: 'NORTH', 2: 'EAST', 3: 'WEST', 4: 'PICKUP',5: 'DROPOFF'}

def save_visual_results(env, doc, fact, cfact, target, statement, cf_typename):
    # create a new document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Menlo'

    orangetext = doc.add_paragraph().add_run('Factual explanation:')
    orangefont = orangetext.font
    orangefont.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    doc.add_paragraph("Passenger location: " + str(fact[2]))
    doc.add_paragraph("Passenger destination: " + str(fact[3]))
    if(len(env.MAP[0]) == 11):
        doc = state_render_doc(env,  doc, fact)
    else:
        doc = state_render_doc8x8(env,  doc, fact)
    if(cfact != None):
        orangetext = doc.add_paragraph().add_run('Counterfactual explanation:')
        orangefont = orangetext.font
        orangefont.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

        doc.add_paragraph("Passenger location: " + str(cfact.cf_state[2]))
        doc.add_paragraph("Passenger destination: " + str(cfact.cf_state[3]))
        doc.add_paragraph("State Path: " + ', '.join(str(i) for i in cfact.path) )
        #doc.add_paragraph("Number of steps: " + str(cfact.num_steps))
        doc.add_paragraph("Loss function value: " + str(cfact.value))
        doc.add_paragraph("Action path: " + ', '.join(str(i) for i in cfact.action_path))
        doc.add_paragraph(statement)

        if(len(env.MAP[0]) == 11):
            doc = state_render_doc(env,  doc, cfact.cf_state, True)
        else:
            doc = state_render_doc8x8(env,  doc, cfact.cf_state, True)
    else:
        redtext = doc.add_paragraph().add_run('No counterfactual found for target ' + ACTIONS[target])
        redfont = redtext.font
        redfont.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    
    # Not at the end of a page, add break
    if(cf_typename != "NETACTION" or ( cf_typename == "NETACTION" and cfact == None)):
        paragraph = doc.paragraphs[-1]
        element = paragraph._element    
        if not(element.get(qn('w:lastRenderedPageBreak')) is not None):
            doc.add_page_break()
    
    return doc

def state_render_doc(env, doc, state, iscf=False):
    
    taxi_row = state[0]
    taxi_col = state[1]
    pass_idx = state[2]
    dest_idx = state[3]
    MAP = env.MAP
    # MAP = [
    #     "+---------+",
    #     "|R: | : :G|",
    #     "| : | : : |",
    #     "| : : : : |",
    #     "| | : | : |",
    #     "|Y| : |B: |",
    #     "+---------+",
    # ]

    #Map row 00
    doc.add_paragraph(list(MAP[0]))

    #Map row 0
    row0 = list(MAP[1])
    add_locrow(doc, row0, taxi_row, taxi_col, pass_idx, dest_idx, 0)

    #Map row 1
    row1 = list(MAP[2])
    if(taxi_row == 1):
        add_taxirow(doc, row1, taxi_col, pass_idx, dest_idx, 1)
    else:
        doc.add_paragraph(''.join(str(i) for i in row1))

    #Map row 2
    row2 = list(MAP[3])
    if(taxi_row == 2):
        add_taxirow(doc, row2, taxi_col, pass_idx, dest_idx, 2)
    else:
        doc.add_paragraph(''.join(str(i) for i in row2))
    
    #Map row 3
    row3 = list(MAP[4])
    if(taxi_row == 3):
        add_taxirow(doc, row3, taxi_col, pass_idx, dest_idx, 3)
    else:
        doc.add_paragraph(''.join(str(i) for i in row3))
    
    #Map row 4
    row4 = list(MAP[5])
    add_locrow(doc, row4, taxi_row, taxi_col, pass_idx, dest_idx, 4)

    #Map row 5
    p = doc.add_paragraph(list(MAP[6]))

    return doc

def add_locrow(doc, txtrow, taxirow, taxicol, pass_idx, dest_idx, rownum):
    GREEN = 'green'
    YELLOW = 'yellow'
    if pass_idx != 4:
        taxi_color = WD_COLOR_INDEX.YELLOW
    else:
        taxi_color = WD_COLOR_INDEX.GREEN
    
    row = doc.add_paragraph()

    if(rownum == 0):
        if pass_idx == 0:
            r_color = RGBColor(0, 0, 255)
        elif dest_idx == 0:
            r_color = RGBColor(255, 0, 255)
        else:
            r_color = RGBColor(0, 0, 0)

        if pass_idx == 1:
            g_color = RGBColor(0, 0, 255)
        elif dest_idx == 1:
            g_color = RGBColor(255, 0, 255)
        else:
            g_color = RGBColor(0, 0, 0)
        
        row.add_run(txtrow[0])
        text = row.add_run('R')
        posfont = text.font
        posfont.color.rgb = r_color
        if(taxicol == 0 and taxirow == 0):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[2])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 1 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[4])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 2 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[6])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 3 and taxirow == 0):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[8])
        
        text = row.add_run('G')
        posfont = text.font
        posfont.color.rgb = g_color
        if(taxicol == 4 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[10])
    elif(rownum == 4):
        if pass_idx == 2:
            y_color = RGBColor(0, 0, 255)
        elif dest_idx == 2:
            y_color = RGBColor(255, 0, 255)
        else:
            y_color = RGBColor(0, 0, 0)

        if pass_idx == 3:
            b_color = RGBColor(0, 0, 255)
        elif dest_idx == 3:
            b_color = RGBColor(255, 0, 255)
        else:
            b_color = RGBColor(0, 0, 0)

        row.add_run(txtrow[0])
        text = row.add_run('Y')
        posfont = text.font
        posfont.color.rgb = y_color
        if(taxicol == 0 and taxirow == 4):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[2])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 1 and taxirow == 4):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[4])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 2 and taxirow == 4):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[6])
        
        text = row.add_run('B')
        posfont = text.font
        posfont.color.rgb = b_color
        if(taxicol == 3 and taxirow == 4):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[8])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 4 and taxirow == 4):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[10])
    else:
        doc.add_paragraph('error')

    return

def add_taxirow(doc, txtrow, taxicol, pass_idx, dest_idx, rownum):
    YELLOW = 'yellow'
    GREEN = 'green'
    if pass_idx != 4:
        taxi_color = WD_COLOR_INDEX.YELLOW
    else:
        taxi_color = WD_COLOR_INDEX.GREEN

    row = doc.add_paragraph()
        
    row.add_run(txtrow[0])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 0):
        posfont.highlight_color = taxi_color
    
    row.add_run(txtrow[2])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 1):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[4])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 2):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[6])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 3):
        posfont.highlight_color = taxi_color
    
    row.add_run(txtrow[8])
    
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 4):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[10])

    return

def state_render_doc8x8(env, doc, state, iscf=False):
    
    taxi_row = state[0]
    taxi_col = state[1]
    pass_idx = state[2]
    dest_idx = state[3]
    MAP = env.MAP
    # MAP = [
    #     "+---------+",
    #     "|R: | : :G|",
    #     "| : | : : |",
    #     "| : : : : |",
    #     "| | : | : |",
    #     "|Y| : |B: |",
    #     "+---------+",
    # ]

    #Map row 00
    doc.add_paragraph(list(MAP[0]))

    #Map row 0
    row0 = list(MAP[1])
    add_locrow8x8(doc, row0, taxi_row, taxi_col, pass_idx, dest_idx, 0)

    #Map row 1
    row1 = list(MAP[2])
    if(taxi_row == 1):
        add_taxirow8x8(doc, row1, taxi_col, pass_idx, dest_idx, 1)
    else:
        doc.add_paragraph(''.join(str(i) for i in row1))

    #Map row 2
    row2 = list(MAP[3])
    if(taxi_row == 2):
        add_taxirow8x8(doc, row2, taxi_col, pass_idx, dest_idx, 2)
    else:
        doc.add_paragraph(''.join(str(i) for i in row2))
    
    #Map row 3
    row3 = list(MAP[4])
    if(taxi_row == 3):
        add_taxirow8x8(doc, row3, taxi_col, pass_idx, dest_idx, 3)
    else:
        doc.add_paragraph(''.join(str(i) for i in row3))
    
    #Map row 4
    row4 = list(MAP[5])
    if(taxi_row == 4):
        add_taxirow8x8(doc, row4, taxi_col, pass_idx, dest_idx, 4)
    else:
        doc.add_paragraph(''.join(str(i) for i in row4))

    #Map row 5
    row5 = list(MAP[6])
    if(taxi_row == 5):
        add_taxirow8x8(doc, row5, taxi_col, pass_idx, dest_idx, 5)
    else:
        doc.add_paragraph(''.join(str(i) for i in row5))

    #Map row 6
    row6 = list(MAP[7])
    if(taxi_row == 6):
        add_taxirow8x8(doc, row6, taxi_col, pass_idx, dest_idx, 6)
    else:
        doc.add_paragraph(''.join(str(i) for i in row6))
    
    #Map row 7
    row7 = list(MAP[8])
    add_locrow8x8(doc, row7, taxi_row, taxi_col, pass_idx, dest_idx, 7)

    #Map row 8
    p = doc.add_paragraph(list(MAP[9]))

    return doc

def add_locrow8x8(doc, txtrow, taxirow, taxicol, pass_idx, dest_idx, rownum):
    GREEN = 'green'
    YELLOW = 'yellow'
    if pass_idx != 4:
        taxi_color = WD_COLOR_INDEX.YELLOW
    else:
        taxi_color = WD_COLOR_INDEX.GREEN
    
    row = doc.add_paragraph()

    if(rownum == 0):
        if pass_idx == 0:
            r_color = RGBColor(0, 0, 255)
        elif dest_idx == 0:
            r_color = RGBColor(255, 0, 255)
        else:
            r_color = RGBColor(0, 0, 0)

        if pass_idx == 1:
            g_color = RGBColor(0, 0, 255)
        elif dest_idx == 1:
            g_color = RGBColor(255, 0, 255)
        else:
            g_color = RGBColor(0, 0, 0)
        
        row.add_run(txtrow[0])
        text = row.add_run('R')
        posfont = text.font
        posfont.color.rgb = r_color
        if(taxicol == 0 and taxirow == 0):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[2])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 1 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[4])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 2 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[6])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 3 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[8])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 4 and taxirow == 0):
            posfont.highlight_color = taxi_color
            
        row.add_run(txtrow[8])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 5 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[12])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 6 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[14])
        text = row.add_run('G')
        posfont = text.font
        posfont.color.rgb = g_color
        if(taxicol == 7 and taxirow == 0):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[16])
    elif(rownum == 7):
        if pass_idx == 2:
            y_color = RGBColor(0, 0, 255)
        elif dest_idx == 2:
            y_color = RGBColor(255, 0, 255)
        else:
            y_color = RGBColor(0, 0, 0)

        if pass_idx == 3:
            b_color = RGBColor(0, 0, 255)
        elif dest_idx == 3:
            b_color = RGBColor(255, 0, 255)
        else:
            b_color = RGBColor(0, 0, 0)

        row.add_run(txtrow[0])
        text = row.add_run('Y')
        posfont = text.font
        posfont.color.rgb = y_color
        if(taxicol == 0 and taxirow == 7):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[2])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 1 and taxirow == 7):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[4])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 2 and taxirow == 7):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[6])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 3 and taxirow == 7):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[8])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 4 and taxirow == 7):
            posfont.highlight_color = taxi_color
        
        row.add_run(txtrow[10])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 5 and taxirow == 7):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[12])
        text = row.add_run('B')
        posfont = text.font
        posfont.color.rgb = b_color
        if(taxicol == 6 and taxirow == 7):
            posfont.highlight_color = taxi_color

        row.add_run(txtrow[14])
        text = row.add_run(' ')
        posfont = text.font
        if(taxicol == 7 and taxirow == 7):
            posfont.highlight_color = taxi_color
 
        row.add_run(txtrow[16])
    else:
        doc.add_paragraph('error'+str(rownum))

    return

def add_taxirow8x8(doc, txtrow, taxicol, pass_idx, dest_idx, rownum):
    YELLOW = 'yellow'
    GREEN = 'green'
    if pass_idx != 4:
        taxi_color = WD_COLOR_INDEX.YELLOW
    else:
        taxi_color = WD_COLOR_INDEX.GREEN

    row = doc.add_paragraph()
        
    row.add_run(txtrow[0])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 0):
        posfont.highlight_color = taxi_color
    
    row.add_run(txtrow[2])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 1):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[4])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 2):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[6])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 3):
        posfont.highlight_color = taxi_color
    
    row.add_run(txtrow[8])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 4):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[10])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 5):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[12])
    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 6):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[14])

    text = row.add_run(' ')
    posfont = text.font
    if(taxicol == 7):
        posfont.highlight_color = taxi_color

    row.add_run(txtrow[16])

    return


